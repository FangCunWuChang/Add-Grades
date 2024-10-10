import os
import pdfplumber
import re
from docx import Document
from docx.shared import Inches
import office
import argparse
import datetime

def extract_grades_from_pdf(pdf_path):
    """从PDF表格中提取姓名和成绩，基于表头匹配"""
    grades_data = []
    column_indices = {}  # 用于保存表头的列索引

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                # 开始处理数据行
                for row in table[0:]:
                    # 假设表头包含"学号", "姓名", "实习报告"等关键词
                    if not column_indices:  # 如果还没有找到表头
                        for i, cell in enumerate(row):  # 假设表头是表格的第一行
                            if "学号" in cell:
                                column_indices["student_id"] = i
                            elif "姓名" in cell:
                                column_indices["name"] = i
                            elif "实习报告" in cell:
                                column_indices["grade"] = i
                        if column_indices:  # 找到表头之后继续处理下一行数据
                            continue

                    try:
                        student_id = row[column_indices["student_id"]].strip()
                        student_name = row[column_indices["name"]].strip()
                        grade = row[column_indices["grade"]].strip()
                        # 验证成绩格式为数字
                        if re.match(r'\d+', grade):
                            grades_data.append((student_id, student_name, grade))
                        else:
                            print("成绩格式错误：", grade)
                    except (IndexError, KeyError):
                        continue  # 跳过可能格式错误的行

    return grades_data

def save_doc_to_docx(path_in):  # doc转docx
    print("转换文件格式：", path_in)
    office.word.doc2docx(input_path=path_in, output_path=os.path.dirname(path_in))
    os.remove(path_in)

def generate_teacher_comment(grade):
    """根据成绩生成教师评语"""
    if 60 <= grade <= 70:
        return "实习报告内容尚可，但需要进一步提高对专业知识的理解。"
    elif 70 < grade <= 85:
        return "实习报告较为完整，体现了较好的专业理解能力。"
    elif 85 < grade <= 100:
        return "实习报告内容优秀，体现了较强的专业素养和实践能力。"
    return ""  # 其他分数段不生成评语

def generate_date_time():
    return datetime.datetime.now().strftime('%Y年%m月%d日')

def fill_grade_in_report(report_path, signature_image_path, name, grade):
    """在实习报告中填写成绩与评语"""
    # 获取绝对路径
    report_path = os.path.abspath(report_path)
    print("处理报告：", report_path)

    try:
        if report_path.endswith(".doc"):
            # 如果是doc文件，先转换为docx
            report_path_docx = report_path.replace('.doc', '.docx')
            save_doc_to_docx(report_path)
            report_path = report_path_docx
        
        # 处理docx文件
        grade_filled = False
        sign_filled = False
        comment_filled = False

        doc = Document(report_path)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if ("综合成绩评定（百分制或五级制）：        " in paragraph.text) and not grade_filled:
                            # 填写成绩
                            paragraph.text = paragraph.text.replace("综合成绩评定（百分制或五级制）：        ", f"综合成绩评定（百分制或五级制）：  {grade}  ")
                            print("已填写成绩")
                            grade_filled = True

                        if ("指导教师手写签名：" in paragraph.text) and not sign_filled:
                            # 加入签名
                            paragraph.text.rstrip()
                            run = paragraph.add_run()
                            run.add_picture(signature_image_path, width=Inches(1.5))  # 调整宽度为1.5英寸
                            print("已插入签名")
                            sign_filled = True
                    
                    if grade_filled:
                        for paragraph in cell.paragraphs:
                            if "（学生是否完成实习计划，实习任务完成的水平、效益，研究和解决实践问题的意识和能力，工作态度、综合素质、品德纪律等情况）" in paragraph.text:
                                paragraph.text = "" # 移除提示信息
                                print("移除提示信息")

                            if "年   月   日" in paragraph.text:
                                paragraph.text = paragraph.text.replace("年   月   日", generate_date_time())
                                print("已填写日期")

                            if (not paragraph.text.strip()) and (not comment_filled):  # 检查段落是否为空
                                paragraph.text = generate_teacher_comment(int(grade)) # 替换为评语
                                print("已插入评语")
                                comment_filled = True

                        break

                if grade_filled:
                        break
            if grade_filled:
                        break

        if grade_filled and sign_filled and comment_filled:
            doc.save(report_path)
            return True

        return False
    
    except Exception as e:
        print(f"Error processing {report_path}: {e}")
        return False

def main(pdf_path, reports_folder, signature_image_path):
    # 提取成绩表中的姓名和成绩
    grades_data = extract_grades_from_pdf(pdf_path)
    print("成绩表：", grades_data)
    
    # 缓存报告文件夹中的文件列表
    report_files = [f for f in os.listdir(reports_folder) if f.endswith(".doc") or f.endswith(".docx")]
    
    unmatched_students = []
    unmatched_reports = []
    failed_reports = []
    processed_files = set()

    # 遍历成绩表中的学生，寻找匹配的报告
    for student_id, student_name, grade in grades_data:
        report_found = False
        for report_filename in report_files:
            report_path = os.path.join(reports_folder, report_filename)
            
            # 检查报告文件名中是否包含该学生姓名
            if student_name in report_filename:
                report_found = True
                success = fill_grade_in_report(report_path, signature_image_path, student_name, grade)

                processed_files.add(report_filename)
                if not success:
                    failed_reports.append(report_filename)
                break  # 找到匹配的报告后不再继续遍历
        
        if not report_found:
            unmatched_students.append((student_id, student_name))

    # 查找未被处理的文件
    for report_filename in report_files:
        if report_filename not in processed_files:
            unmatched_reports.append(report_filename)

    # 输出未匹配的学生或报告
    if unmatched_students:
        print("未匹配的学生信息：", unmatched_students)
    if unmatched_reports:
        print("未处理的实习报告：", unmatched_reports)
    if failed_reports:
        print("处理失败的实习报告：", failed_reports)

# 运行主程序
if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="将学生成绩从成绩表填入实习报告")
    parser.add_argument("--directory", type=str, help="实习报告目录")
    parser.add_argument("--grades", type=str, help="成绩表路径")
    parser.add_argument("--signature", type=str, help="签名图片路径")

    args = parser.parse_args()
    
    main(args.grades, args.directory, args.signature)
